"""
file_extractor.py — Binary → Text Extraction Utility
====================================================
Converts binary file content (PDF, DOCX, TXT, CSV) into plain text
for Presidio analysis. Handles corrupted/encrypted files gracefully.
"""

import io
import csv
import logging
from typing import Optional

logger = logging.getLogger("scanner.file_extractor")


def extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyPDF2."""
    try:
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))

        # Detect encrypted PDFs
        if reader.is_encrypted:
            logger.warning("📄 PDF is encrypted — attempting empty-password decrypt")
            try:
                reader.decrypt("")
            except Exception:
                logger.warning("📄 PDF decryption failed — skipping file")
                return ""

        texts = []
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    texts.append(text)
            except Exception as e:
                logger.debug(f"  Page {page_num + 1} extraction error: {e}")

        result = "\n".join(texts)
        logger.debug(f"  PDF extracted {len(result)} chars from {len(reader.pages)} pages")
        return result

    except Exception as e:
        logger.error(f"❌ PDF extraction failed: {e}")
        return ""


def extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx (paragraphs + tables)."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Table cells
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))

        # Headers and footers
        for section in doc.sections:
            for header in [section.header, section.footer]:
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)

        result = "\n".join(parts)
        logger.debug(f"  DOCX extracted {len(result)} chars")
        return result

    except Exception as e:
        logger.error(f"❌ DOCX extraction failed: {e}")
        return ""


def extract_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text bytes with multi-encoding fallback."""
    # Try common encodings in order of preference
    encodings = ["utf-8", "utf-8-sig", "euc-kr", "cp949", "latin-1"]

    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            logger.debug(f"  TXT decoded with '{encoding}': {len(text)} chars")
            return text
        except (UnicodeDecodeError, LookupError):
            continue

    logger.warning("  TXT: all encodings failed — using lossy UTF-8 decode")
    return file_bytes.decode("utf-8", errors="replace")


def extract_from_csv(file_bytes: bytes) -> str:
    """Extract all cell text from CSV bytes."""
    try:
        text = extract_from_txt(file_bytes)
        reader = csv.reader(io.StringIO(text))
        parts = []
        for row in reader:
            parts.append(" ".join(cell.strip() for cell in row if cell.strip()))
        result = "\n".join(parts)
        logger.debug(f"  CSV extracted {len(result)} chars")
        return result
    except Exception as e:
        logger.error(f"❌ CSV extraction failed: {e}")
        return ""


# ── MIME Type Dispatcher ───────────────────────────────────────────────────────

_EXTRACTOR_MAP = {
    "application/pdf": extract_from_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_from_docx,
    "text/plain": extract_from_txt,
    "text/csv": extract_from_csv,
    # Google native export types
    "application/vnd.google-apps.document": extract_from_txt,
    "application/vnd.google-apps.spreadsheet": extract_from_csv,
}


def extract_text(file_bytes: bytes, mime_type: str) -> Optional[str]:
    """
    Main dispatcher: extract plain text from file bytes based on MIME type.

    Args:
        file_bytes: Raw binary content of the file.
        mime_type:  MIME type string (e.g., 'application/pdf').

    Returns:
        Extracted text string, or None if unsupported/failed.
    """
    if not file_bytes:
        logger.debug("  Empty file bytes — skipping")
        return None

    extractor = _EXTRACTOR_MAP.get(mime_type)
    if extractor is None:
        logger.debug(f"  Unsupported MIME type: {mime_type}")
        return None

    try:
        text = extractor(file_bytes)
        return text if text.strip() else None
    except Exception as e:
        logger.error(f"❌ Unexpected extraction error for {mime_type}: {e}")
        return None
